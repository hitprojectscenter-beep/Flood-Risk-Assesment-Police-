# -*- coding: utf-8 -*-
"""
Generate 4 specification update documents:
- BRD_Updates_Hebrew.docx (Business Requirements appendix - HE)
- BRD_Updates_English.docx (Business Requirements appendix - EN)
- TDD_Updates_Hebrew.docx (Technical Design appendix - HE)
- TDD_Updates_English.docx (Technical Design appendix - EN)

These are appendices to the original BRD/TDD that document all changes
made during development sessions 1-N (Israel weights, sliders, i18n, etc.)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

NAVY = RGBColor(0x0F, 0x17, 0x2A)
TEAL = RGBColor(0x14, 0xB8, 0xA6)
DARK_GRAY = RGBColor(0x33, 0x41, 0x55)
GRAY = RGBColor(0x64, 0x74, 0x8B)

OUTPUT_DIR = r'C:\Users\imark\Desktop\משימת מבחן\files_2'


def setup_doc(rtl=True):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    if rtl:
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Arial'
        if rtl:
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.styles['Heading 1'].font.size = Pt(20)
    doc.styles['Heading 1'].font.color.rgb = TEAL
    doc.styles['Heading 2'].font.size = Pt(15)
    doc.styles['Heading 2'].font.color.rgb = NAVY
    doc.styles['Heading 3'].font.size = Pt(12)
    doc.styles['Heading 3'].font.color.rgb = DARK_GRAY
    return doc


def add_para(doc, text, rtl=True, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    p.alignment = align or (WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT)
    if rtl:
        pPr = p._element.get_or_add_pPr()
        bidi = pPr.makeelement(qn('w:bidi'), {})
        pPr.append(bidi)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1, rtl=True):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        pPr = h._element.get_or_add_pPr()
        bidi = pPr.makeelement(qn('w:bidi'), {})
        pPr.append(bidi)
    return h


def add_req(doc, code, title, desc, rtl=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        pPr = p._element.get_or_add_pPr()
        bidi = pPr.makeelement(qn('w:bidi'), {})
        pPr.append(bidi)
    code_run = p.add_run(f'{code}  ')
    code_run.bold = True
    code_run.font.color.rgb = TEAL
    code_run.font.name = 'Arial'
    code_run.font.size = Pt(11)
    title_run = p.add_run(f'{title}\n')
    title_run.bold = True
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(11)
    desc_run = p.add_run(desc)
    desc_run.font.name = 'Arial'
    desc_run.font.size = Pt(10)
    desc_run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(10)


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.color.rgb = TEAL
    run.font.size = Pt(8)


# ============================================================
# BRD HEBREW
# ============================================================
def create_brd_hebrew():
    doc = setup_doc(rtl=True)

    add_heading(doc, 'TSRS — נספח עדכוני BRD', level=1)
    add_para(doc, 'מסמך אפיון דרישות עסקיות — נספח עדכונים', size=13, color=GRAY)
    add_para(doc, 'גרסה 2.0 | אפריל 2026', size=11, color=GRAY)
    add_separator(doc)

    add_heading(doc, '1. מבוא', level=2)
    add_para(doc, 'נספח זה מתעד את כל העדכונים שבוצעו ל-BRD המקורי במהלך פיתוח מערכת TSRS. '
                  'העדכונים נובעים מתוצאות הפיתוח, משוב מהמשתמשים, וכיול המודל לתנאי ישראל 2026.')

    add_heading(doc, '2. עדכוני דרישות עסקיות', level=2)

    add_heading(doc, '2.1 משקלי מודל TSRS', level=3)
    add_req(doc, 'REQ-W01', 'כיול משקלי TSRS למדינת ישראל',
            'המשקלים המקוריים (35/30/20/10/5) הותאמו לתנאי ישראל 2026 ל-(25/30/15/18/12). '
            'הכיול מבוסס על: (1) רצועת חוף צרה — חשיפה אחידה בין ערים, (2) פערי משאבים '
            'גדולים בין תחנות מרכז ופריפריה, (3) מקלטים אנכיים מוגבלים בערים קטנות.')

    add_req(doc, 'REQ-W02', 'סליידרים אינטראקטיביים לכוונון משקלים',
            'המערכת תאפשר למשתמש לכוונן את 5 משקלי TSRS בזמן אמת באמצעות סליידרים. '
            'תוצג נוסחת TSRS דינמית. הציונים על המפה יתעדכנו מיידית. שני פריסטים: '
            '🇮🇱 ישראל (ברירת מחדל) | 🌍 בינלאומי. נרמול אוטומטי לסכום 100%.')

    add_heading(doc, '2.2 נתונים', level=3)
    add_req(doc, 'REQ-D01', 'נתוני סוציו-אקונומיים אמיתיים מלמ"ס',
            'המערכת תשתמש בנתוני אשכול סוציו-אקונומי אמיתיים מ-202 ערים בישראל '
            '(הלשכה המרכזית לסטטיסטיקה, 2022). הנתונים מקובץ "פרופיל חברתי כלכלי לפי ישובים".')

    add_req(doc, 'REQ-D02', 'גבולות מוניציפליים מ-OpenStreetMap',
            'המערכת תציג את הגבולות המוניציפליים האמיתיים של 554 רשויות בישראל, '
            'מ-OSM admin_level 7+8. הסינון יוצג רק לערים עם נתוני CBS זמינים.')

    add_heading(doc, '2.3 ממשק משתמש', level=3)
    add_req(doc, 'REQ-L01', 'תמיכה ב-5 שפות',
            'המערכת תתמוך ב-5 שפות: עברית (ברירת מחדל, RTL), אנגלית, רוסית, צרפתית, '
            'וספרדית. בורר שפות בכותרת. שמירת העדפה ב-localStorage.')

    add_req(doc, 'REQ-M01', 'התאמה למובייל וטאבלט',
            'המערכת תתמוך ב-5 נקודות שבירה רספונסיביות: 1023px (טאבלט), 767px '
            '(טאבלט קטן עם סרגל צד מתקפל), 480px (טלפון), 360px (טלפון קטן), '
            'ו-600px (פריסת מודל). מטרות מגע גדולות על מובייל.')

    add_req(doc, 'REQ-I01', 'בועית מידע אוטומטית',
            'בכניסה ליישום תיפתח אוטומטית בועית מידע מפורטת על המודל TSRS. '
            'תכלול: בעיה, נוסחה, 5 מדדים, יישום, ביבליוגרפיה. כפתור "הסבר" '
            'בכותרת לפתיחה מחדש.')

    add_heading(doc, '2.4 ויזואליזציה', level=3)
    add_req(doc, 'REQ-V01', 'מפת חום הצפות דינמית',
            'אזורי ההצפה יוצגו כמפת חום (heatmap) דינמית המשתנה עם סליידר גובה הגל. '
            'גרדיאנט: כחול → ירוק → צהוב → כתום → אדום. רדיוס מתאים לעוצמה.')

    add_req(doc, 'REQ-V02', 'פילוח גאוגרפי בלחיצה על עיר',
            'בלחיצה על עיר תוצג בסרגל הצד פילוח דמוגרפי כולל: התפלגות גילאים '
            '(גרף עמודות), אשכול ס"א, הכנסה ממוצעת, ובעלות רכב.')

    add_heading(doc, '2.5 שכבות מפה', level=3)
    add_req(doc, 'REQ-O01', 'שכבת תחנות משטרה אמיתית מ-OSM',
            'נקודות תחנות המשטרה ייטענו דינמית מ-OSM Overpass API '
            '(amenity=police). יוצגו עם אייקון משטרת ישראל. זמין מזום 12+.')

    add_req(doc, 'REQ-O02', 'סימבולוגיית מבנים לפי עומק הצפה',
            'מבנים מ-OSM ייצבעו לפי גובהם ביחס לעומק ההצפה: אדום (מתחת), '
            'ירוק (מעל), כחול (מקלט פוטנציאלי 4+ קומות). זמין מזום 15+.')

    add_heading(doc, '3. דרישות שהוסרו', level=2)
    add_para(doc, '• שכבת קו חוף — הוסרה (נתונים לא מדויקים)', size=11)
    add_para(doc, '• דרישה ל-FastAPI backend — הוחלפה בארכיטקטורה סטטית (Vercel)', size=11)

    add_heading(doc, '4. הערות', level=2)
    add_para(doc, 'נספח זה אינו מחליף את ה-BRD המקורי, אלא מוסיף ומעדכן אותו. '
                  'במקרה של סתירה, נספח זה גובר.')
    add_para(doc, 'מקום שמירה: tsrs-app/files_2/BRD_Updates_Hebrew.docx', size=10, color=GRAY)

    doc.save(os.path.join(OUTPUT_DIR, 'BRD_Updates_Hebrew.docx'))
    print('OK BRD_Updates_Hebrew.docx')


# ============================================================
# BRD ENGLISH
# ============================================================
def create_brd_english():
    doc = setup_doc(rtl=False)

    add_heading(doc, 'TSRS — BRD Updates Appendix', level=1, rtl=False)
    add_para(doc, 'Business Requirements Document — Updates Appendix', rtl=False, size=13, color=GRAY)
    add_para(doc, 'Version 2.0 | April 2026', rtl=False, size=11, color=GRAY)
    add_separator(doc)

    add_heading(doc, '1. Introduction', level=2, rtl=False)
    add_para(doc, 'This appendix documents all updates made to the original BRD during '
                  'TSRS system development. The updates result from development outcomes, '
                  'user feedback, and model calibration for Israel 2026 conditions.', rtl=False)

    add_heading(doc, '2. Business Requirement Updates', level=2, rtl=False)

    add_heading(doc, '2.1 TSRS Model Weights', level=3, rtl=False)
    add_req(doc, 'REQ-W01', 'Israel-Calibrated TSRS Weights',
            'Original weights (35/30/20/10/5) have been calibrated to Israel 2026 '
            'conditions: (25/30/15/18/12). Calibration based on: (1) narrow Mediterranean '
            'coastline — uniform exposure across cities, (2) large resource gaps between '
            'central and peripheral stations, (3) limited vertical shelters in small towns.', rtl=False)

    add_req(doc, 'REQ-W02', 'Interactive Weight Adjustment Sliders',
            'The system shall allow users to adjust the 5 TSRS weights in real-time via '
            'sliders. Dynamic TSRS formula display. Map scores update instantly. Two presets: '
            '🇮🇱 Israel (default) | 🌍 International. Auto-normalization to 100% sum.', rtl=False)

    add_heading(doc, '2.2 Data', level=3, rtl=False)
    add_req(doc, 'REQ-D01', 'Real CBS Socioeconomic Data',
            'The system shall use real socioeconomic cluster data from 202 Israeli cities '
            '(Central Bureau of Statistics, 2022). Source: "Socioeconomic Profile by '
            'Settlements" file.', rtl=False)

    add_req(doc, 'REQ-D02', 'Municipal Boundaries from OpenStreetMap',
            'The system shall display real municipal boundaries of 554 Israeli '
            'authorities, from OSM admin_level 7+8. Filtering: only cities with '
            'available CBS data are shown.', rtl=False)

    add_heading(doc, '2.3 User Interface', level=3, rtl=False)
    add_req(doc, 'REQ-L01', '5-Language Support',
            'The system shall support 5 languages: Hebrew (default, RTL), English, Russian, '
            'French, and Spanish. Language selector in header. Preference saved to localStorage.', rtl=False)

    add_req(doc, 'REQ-M01', 'Mobile and Tablet Responsive Design',
            'The system shall support 5 responsive breakpoints: 1023px (tablet), 767px '
            '(small tablet with collapsible sidebar), 480px (phone), 360px (small phone), '
            '600px (modal layout). Larger touch targets on mobile.', rtl=False)

    add_req(doc, 'REQ-I01', 'Auto-Open Info Modal',
            'On app entry, an info modal shall auto-open with detailed TSRS model '
            'explanation including: problem, formula, 5 metrics, application, bibliography. '
            '"Info" button in header to reopen.', rtl=False)

    add_heading(doc, '2.4 Visualization', level=3, rtl=False)
    add_req(doc, 'REQ-V01', 'Dynamic Inundation Heatmap',
            'Flood zones shall be displayed as a dynamic heatmap that changes with the '
            'wave height slider. Gradient: blue → green → yellow → orange → red. '
            'Radius adapts to intensity.', rtl=False)

    add_req(doc, 'REQ-V02', 'Geographic Profile on City Click',
            'On city click, the sidebar shall display a demographic profile including: '
            'age distribution (bar chart), socioeconomic cluster, average income, and '
            'car ownership.', rtl=False)

    add_heading(doc, '2.5 Map Layers', level=3, rtl=False)
    add_req(doc, 'REQ-O01', 'Real OSM Police Stations Layer',
            'Police station points shall be loaded dynamically from OSM Overpass API '
            '(amenity=police). Displayed with Israel Police icon. Available at zoom 12+.', rtl=False)

    add_req(doc, 'REQ-O02', 'Building Symbology by Flood Depth',
            'OSM buildings shall be colored by their height relative to flood depth: '
            'red (below), green (above), blue (potential shelter, 4+ floors). '
            'Available at zoom 15+.', rtl=False)

    add_heading(doc, '3. Removed Requirements', level=2, rtl=False)
    add_para(doc, '• Coastline layer — removed (inaccurate data)', rtl=False, size=11)
    add_para(doc, '• FastAPI backend requirement — replaced with static architecture (Vercel)', rtl=False, size=11)

    add_heading(doc, '4. Notes', level=2, rtl=False)
    add_para(doc, 'This appendix does not replace the original BRD but supplements and '
                  'updates it. In case of conflict, this appendix takes precedence.', rtl=False)
    add_para(doc, 'Storage: tsrs-app/files_2/BRD_Updates_English.docx', rtl=False, size=10, color=GRAY)

    doc.save(os.path.join(OUTPUT_DIR, 'BRD_Updates_English.docx'))
    print('OK BRD_Updates_English.docx')


# ============================================================
# TDD HEBREW
# ============================================================
def create_tdd_hebrew():
    doc = setup_doc(rtl=True)

    add_heading(doc, 'TSRS — נספח עדכוני TDD', level=1)
    add_para(doc, 'מסמך אפיון טכני — נספח עדכונים', size=13, color=GRAY)
    add_para(doc, 'גרסה 2.0 | אפריל 2026', size=11, color=GRAY)
    add_separator(doc)

    add_heading(doc, '1. עדכוני ארכיטקטורה', level=2)
    add_para(doc, 'הארכיטקטורה המקורית (React + FastAPI + PostGIS) הוחלפה בארכיטקטורה '
                  'סטטית קלת-משקל המתאימה לפריסה ב-Vercel ללא backend.', bold=True)

    add_heading(doc, '1.1 ארכיטקטורה חדשה', level=3)
    add_para(doc, '• Frontend: HTML5 + JavaScript ES6 + Leaflet.js (ללא Build step)', size=11)
    add_para(doc, '• Storage: קבצי JSON סטטיים (cities.json, socioeconomic.json, inundation.json)', size=11)
    add_para(doc, '• Map: Leaflet 1.9 + Leaflet.heat + osmtogeojson', size=11)
    add_para(doc, '• Hosting: Vercel (סטטי, חינם)', size=11)
    add_para(doc, '• Backup: Firebase Realtime Database (אופציונלי)', size=11)
    add_para(doc, '• Backend (אופציונלי): Python FastAPI לפיתוח מקומי', size=11)

    add_heading(doc, '2. מודולים חדשים', level=2)

    add_heading(doc, '2.1 frontend/js/weights.js', level=3)
    add_para(doc, 'מודול ניהול משקלי TSRS עם פריסט ישראלי וברירת מחדל בינלאומית. '
                  'API: get(), set(), setAll(), resetToIsrael(), resetToInternational(), '
                  'calculate(H,V,O,R,I), getTier(score), getFormulaString(), onChange().')

    add_heading(doc, '2.2 frontend/js/i18n.js', level=3)
    add_para(doc, 'מודול תרגום לחמש שפות (HE, EN, RU, FR, ES). ~80 מפתחות תרגום לכל שפה. '
                  'API: t(key, params), setLanguage(lang), getLang(), isRTL(), init(). '
                  'שמירת העדפה ב-localStorage.')

    add_heading(doc, '2.3 frontend/js/osm-overlays.js', level=3)
    add_para(doc, 'אינטגרציה עם Overpass API לטעינה דינמית של שכבות OSM: כבישים (זום 13+), '
                  'מבנים (זום 15+), תחנות משטרה (זום 12+). Debouncing, AbortController, '
                  'fallback בין שני endpoint.')

    add_heading(doc, '2.4 frontend/js/inundation.js', level=3)
    add_para(doc, 'שכבת מפת חום הצפות באמצעות Leaflet.heat. גרדיאנט 6 צבעים. '
                  'רדיוס דינמי לפי גובה גל (18-35px). Inland gradient factor.')

    add_heading(doc, '3. נתונים', level=2)
    add_para(doc, '• frontend/data/cities.json — 237 ערים, 3.1MB, גבולות מ-OSM', size=11)
    add_para(doc, '• frontend/data/socioeconomic.json — 202 ערים, 19.6KB, אשכול ס"א מלמ"ס', size=11)
    add_para(doc, '• frontend/data/inundation.json — 20 רמות גובה גל (0.5-10m)', size=11)
    add_para(doc, '• frontend/data/stations.json — נתוני תחנות סינתטיים (גיבוי)', size=11)

    add_heading(doc, '4. CSS Responsive', level=2)
    add_para(doc, '5 נקודות שבירה:', bold=True)
    add_para(doc, '• ≤1023px: טאבלט — סרגל צד 270px, פונטים קטנים', size=11)
    add_para(doc, '• ≤767px: טאבלט קטן — סרגל צד מתקפל overlay, כפתור המבורגר', size=11)
    add_para(doc, '• ≤600px: מודל — grid חד-עמודי', size=11)
    add_para(doc, '• ≤480px: טלפון — כותרת 50px, מטרות מגע 22px slider', size=11)
    add_para(doc, '• ≤360px: טלפון קטן — מינימליסטי, סליידרים קומפקטיים', size=11)

    add_heading(doc, '5. נוסחת חישוב TSRS — משקלים מכוילים', level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TSRS = (H × 0.25) + (V × 0.30) + (O × 0.15) + (R⁻¹ × 0.18) + (I⁻¹ × 0.12)')
    run.font.name = 'Courier New'
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL
    run.bold = True

    add_heading(doc, '6. הערות', level=2)
    add_para(doc, 'נספח זה משלים את ה-TDD המקורי. הוא משקף את המצב הנוכחי של המערכת '
                  'לאחר 15+ commits ב-GitHub.')
    add_para(doc, 'Repository: github.com/hitprojectscenter-beep/Flood-Risk-Assesment-Police-', size=10, color=GRAY)

    doc.save(os.path.join(OUTPUT_DIR, 'TDD_Updates_Hebrew.docx'))
    print('OK TDD_Updates_Hebrew.docx')


# ============================================================
# TDD ENGLISH
# ============================================================
def create_tdd_english():
    doc = setup_doc(rtl=False)

    add_heading(doc, 'TSRS — TDD Updates Appendix', level=1, rtl=False)
    add_para(doc, 'Technical Design Document — Updates Appendix', rtl=False, size=13, color=GRAY)
    add_para(doc, 'Version 2.0 | April 2026', rtl=False, size=11, color=GRAY)
    add_separator(doc)

    add_heading(doc, '1. Architecture Updates', level=2, rtl=False)
    add_para(doc, 'The original architecture (React + FastAPI + PostGIS) has been replaced '
                  'with a lightweight static architecture suitable for Vercel deployment '
                  'without a backend.', rtl=False, bold=True)

    add_heading(doc, '1.1 New Architecture', level=3, rtl=False)
    add_para(doc, '• Frontend: HTML5 + JavaScript ES6 + Leaflet.js (no build step)', rtl=False, size=11)
    add_para(doc, '• Storage: Static JSON files (cities.json, socioeconomic.json, inundation.json)', rtl=False, size=11)
    add_para(doc, '• Map: Leaflet 1.9 + Leaflet.heat + osmtogeojson', rtl=False, size=11)
    add_para(doc, '• Hosting: Vercel (static, free tier)', rtl=False, size=11)
    add_para(doc, '• Backup: Firebase Realtime Database (optional)', rtl=False, size=11)
    add_para(doc, '• Backend (optional): Python FastAPI for local development', rtl=False, size=11)

    add_heading(doc, '2. New Modules', level=2, rtl=False)

    add_heading(doc, '2.1 frontend/js/weights.js', level=3, rtl=False)
    add_para(doc, 'TSRS weight management module with Israel preset and international default. '
                  'API: get(), set(), setAll(), resetToIsrael(), resetToInternational(), '
                  'calculate(H,V,O,R,I), getTier(score), getFormulaString(), onChange().', rtl=False)

    add_heading(doc, '2.2 frontend/js/i18n.js', level=3, rtl=False)
    add_para(doc, '5-language translation module (HE, EN, RU, FR, ES). ~80 translation keys '
                  'per language. API: t(key, params), setLanguage(lang), getLang(), isRTL(), '
                  'init(). Preference saved to localStorage.', rtl=False)

    add_heading(doc, '2.3 frontend/js/osm-overlays.js', level=3, rtl=False)
    add_para(doc, 'Overpass API integration for dynamic OSM layer loading: roads (zoom 13+), '
                  'buildings (zoom 15+), police stations (zoom 12+). Debouncing, '
                  'AbortController, fallback between two endpoints.', rtl=False)

    add_heading(doc, '2.4 frontend/js/inundation.js', level=3, rtl=False)
    add_para(doc, 'Inundation heatmap layer using Leaflet.heat. 6-color gradient. '
                  'Dynamic radius based on wave height (18-35px). Inland gradient factor.', rtl=False)

    add_heading(doc, '3. Data', level=2, rtl=False)
    add_para(doc, '• frontend/data/cities.json — 237 cities, 3.1MB, OSM boundaries', rtl=False, size=11)
    add_para(doc, '• frontend/data/socioeconomic.json — 202 cities, 19.6KB, CBS clusters', rtl=False, size=11)
    add_para(doc, '• frontend/data/inundation.json — 20 wave height levels (0.5-10m)', rtl=False, size=11)
    add_para(doc, '• frontend/data/stations.json — synthetic station data (fallback)', rtl=False, size=11)

    add_heading(doc, '4. Responsive CSS', level=2, rtl=False)
    add_para(doc, '5 breakpoints:', rtl=False, bold=True)
    add_para(doc, '• ≤1023px: tablet — sidebar 270px, smaller fonts', rtl=False, size=11)
    add_para(doc, '• ≤767px: small tablet — collapsible sidebar overlay, hamburger button', rtl=False, size=11)
    add_para(doc, '• ≤600px: modal — single-column grid', rtl=False, size=11)
    add_para(doc, '• ≤480px: phone — 50px header, 22px slider touch targets', rtl=False, size=11)
    add_para(doc, '• ≤360px: small phone — minimalist, compact sliders', rtl=False, size=11)

    add_heading(doc, '5. TSRS Calculation Formula — Calibrated Weights', level=2, rtl=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TSRS = (H × 0.25) + (V × 0.30) + (O × 0.15) + (R⁻¹ × 0.18) + (I⁻¹ × 0.12)')
    run.font.name = 'Courier New'
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL
    run.bold = True

    add_heading(doc, '6. Notes', level=2, rtl=False)
    add_para(doc, 'This appendix supplements the original TDD. It reflects the current '
                  'state of the system after 15+ commits on GitHub.', rtl=False)
    add_para(doc, 'Repository: github.com/hitprojectscenter-beep/Flood-Risk-Assesment-Police-',
             rtl=False, size=10, color=GRAY)

    doc.save(os.path.join(OUTPUT_DIR, 'TDD_Updates_English.docx'))
    print('OK TDD_Updates_English.docx')


if __name__ == '__main__':
    print('Generating BRD/TDD Updates...')
    create_brd_hebrew()
    create_brd_english()
    create_tdd_hebrew()
    create_tdd_english()
    print('Done. 4 files created in', OUTPUT_DIR)
