# -*- coding: utf-8 -*-
"""
Create TSRS Presenter Guide — detailed Word document for each slide.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# === Page Setup (A4, RTL-friendly margins) ===
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# === Style Setup ===
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
style.paragraph_format.space_after = Pt(6)
# Set RTL for normal style
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Arial'
    heading_style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 1'].font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 2'].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
doc.styles['Heading 3'].font.size = Pt(13)

def add_rtl_paragraph(text, style='Normal', bold=False, color=None, size=None, space_after=None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Set bidi (RTL) on paragraph
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)
    run = p.add_run(text)
    run.font.name = 'Arial'
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading_rtl(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = h._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)
    return h

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_key_point(label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)
    run1 = p.add_run(f'{label}: ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
    run1.font.size = Pt(11)
    run1.font.name = 'Arial'
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(4)

# ================================================
# COVER PAGE
# ================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TSRS')
run.font.size = Pt(48)
run.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
run.bold = True
run.font.name = 'Arial'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tsunami Station Risk Score')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
run.font.name = 'Arial'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('מדריך מציג — הסבר מפורט לכל שקף')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
run.bold = True
run.font.name = 'Arial'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('משטרת ישראל — אגף תכנון וניהול מבצעי')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run.font.name = 'Arial'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('יוני 2025')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run.font.name = 'Arial'

doc.add_page_break()

# ================================================
# TABLE OF CONTENTS
# ================================================
add_heading_rtl('תוכן עניינים', level=1)

toc_items = [
    ('שקף 1', 'שער — TSRS Tsunami Station Risk Score'),
    ('שקף 2', 'הבעיה והחזון — ניהול אינטואיטיבי בסביבה קריטית'),
    ('שקף 3', 'המודל — 5 מדדים משוקללים'),
    ('שקף 4', 'דיאגרמה — קלט, עיבוד, פלט'),
    ('שקף 5', 'היישום — GIS אינטראקטיבי'),
    ('שקף 6', 'מקורות מידע וביבליוגרפיה'),
    ('שקף 7', 'סיכום וקריאה לפעולה'),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)
    run1 = p.add_run(f'{num}  ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
    run1.font.size = Pt(12)
    run1.font.name = 'Arial'
    run2 = p.add_run(title)
    run2.font.size = Pt(12)
    run2.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ================================================
# SLIDE 1: Title
# ================================================
add_heading_rtl('שקף 1 — שער', level=1)
add_heading_rtl('TSRS — Tsunami Station Risk Score', level=2)

add_rtl_paragraph('מטרת השקף:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'שקף הפתיחה מציג את שם המודל ואת הקשר הארגוני. המטרה היא ליצור רושם ראשוני '
    'מקצועי ולהבהיר שמדובר במודל גאוגרפי מבוסס מחקר לתיעדוף הקצאת משאבים '
    'בתרחישי צונאמי עבור משטרת ישראל.'
)

add_rtl_paragraph('מה להדגיש בהצגה:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'שם המודל TSRS (Tsunami Station Risk Score) — ציון מספרי (0-100) '
    'שמאפשר דירוג ותיעדוף של הקצאת משאבים בין תחנות משטרה לפי רמת הסיכון מצונאמי.'
)
add_rtl_paragraph(
    'המודל מבוסס על מחקרים אקדמיים מוכרים: Wood & Jones (2015) מ-USGS '
    'להערכת חשיפה לצונאמי, Tucson Workload Model להקצאת כוח אדם משטרתי, '
    'ומחקרי ABM (Agent-Based Modeling) לסימולציית פינוי אוכלוסייה.'
)

add_rtl_paragraph('הקשר:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'המודל פותח לשימוש אגף תכנון וניהול מבצעי של משטרת ישראל, כחלק מההיערכות '
    'למצבי חירום לאומיים הכוללים אסונות טבע. המשטרה היא הגורם הראשון להגיע '
    'לזירה (First Responder) ולכן זקוקה לכלי מבוסס נתונים לקבלת החלטות מהירה.'
)

add_separator()

# ================================================
# SLIDE 2: Problem & Vision
# ================================================
add_heading_rtl('שקף 2 — הבעיה והחזון', level=1)
add_heading_rtl('ניהול אינטואיטיבי בסביבה קריטית', level=2)

add_rtl_paragraph('מטרת השקף:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'להציג את הפער הקיים — העדר כלי GIS ייעודי בתחנות המשטרה — '
    'ואת החזון למערכת שתאפשר קבלת החלטות מבוססת נתונים.'
)

add_rtl_paragraph('ארבע הבעיות המרכזיות:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)

problems = [
    ('פער מידע', 'אין מפת הצפות ייעודית המשולבת עם טריטוריות של תחנות משטרה. כיום, אין למפקד תחנה כלי מרחבי שמראה איזה חלק מהשטח שלו נמצא בסיכון הצפה.'),
    ('הקצאה אינטואיטיבית', 'כוחות משטרה מוקצים למשימות חירום על בסיס ניסיון ואינטואיציה, ללא שיקלול כמותי של צפיפות אוכלוסייה, שיעור קשישים, מוגבלות ניידות, או מאפיינים סוציו-אקונומיים.'),
    ('זמן התראה קצר', 'הים התיכון הוא סביבת Near-field — זמן ההתראה מרגע רעידת האדמה ועד הגעת הגל הוא פחות מ-20 דקות. אין מקום לאלתור או לתכנון בזמן אמת.'),
    ('משימות המשטרה באסון', 'המשטרה אחראית על: (1) הסדרת וניהול פינוי אוכלוסייה לפי פקודה מוכנה מראש, (2) ניהול תנועה מ-ו-אל זירות הפעילות, (3) סגירת זירות ומניעת ביזה.'),
]

for title, desc in problems:
    add_key_point(title, desc)

add_rtl_paragraph('')
add_rtl_paragraph('החזון:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'מפקד תחנה פותח את המערכת, בוחר גובה גל ונפה, ורואה מיידית: '
    'מה הסיכון בטריטוריה שלו, כמה ניידות גיבוי הוא צריך, ואיפה להציב אותן. '
    'המערכת מספקת הנחיות מבצעיות מוכנות כולל צירי פינוי, מקלטים אנכיים, '
    'ונקודות קריטיות לחסימת תנועה.'
)

add_rtl_paragraph('')
add_rtl_paragraph('הנחות בסיס (כולל הנחות סמויות):', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)

assumptions = [
    'המשטרה היא ה-First Responder — פועלת לפני פיקוד העורף ולפני מל"ח',
    'נתוני DEM (מודל גובה) ונתוני אוכלוסייה זמינים ומעודכנים מלמ"ס ומפ"י',
    'מודל ההצפה מבוסס על גובה גל בלבד — לא על מיקום מקור הרעידה או בתימטריה',
    'סימולציית ABM מבוססת על הנחות תנועה ממוצעות ולא על נתוני תנועה בזמן אמת',
    'אשכול סוציו-אקונומי של הלמ"ס משמש כ-proxy לניידות ופגיעות האוכלוסייה',
    'הנחה סמויה: גבולות מוניציפליים תואמים בקירוב לטריטוריות תחנות משטרה',
]
for a in assumptions:
    add_rtl_paragraph(f'   •  {a}', size=11)

add_separator()

# ================================================
# SLIDE 3: Model Components
# ================================================
add_heading_rtl('שקף 3 — המודל: 5 מדדים משוקללים', level=1)
add_heading_rtl('TSRS = (H × 0.35) + (V × 0.30) + (O × 0.20) + (R⁻¹ × 0.10) + (I⁻¹ × 0.05)', level=2)

add_rtl_paragraph('מטרת השקף:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'להציג את הנוסחה המתמטית של המודל ואת חמשת המדדים המרכיבים אותו. '
    'כל מדד מנורמל לסקאלה של 0-100 (Min-Max Normalization) '
    'ומשוקלל לפי חשיבותו היחסית.'
)

add_rtl_paragraph('')
add_rtl_paragraph('פירוט 5 המדדים:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)

components = [
    ('H — חשיפה פיזית (Hazard)', '35%',
     'מודד את רמת החשיפה הפיזית של האזור לצונאמי. מחושב על בסיס: קרבה לקו החוף (מרחק בק"מ), '
     'אחוז השטח המוצף בתרחיש נתון (לפי DEM ומודל הצפה), וזמן הגעת הגל. '
     'ככל שהאזור קרוב יותר לחוף ומשטחו חשוף יותר — הציון גבוה יותר.'),

    ('V — פגיעות אוכלוסייה (Vulnerability)', '30%',
     'מודד את פגיעות האוכלוסייה המתגוררת באזור. מחושב מנתוני למ"ס: '
     'שיעור קשישים (65+), שיעור ילדים (0-14), אשכול סוציו-אקונומי (כ-proxy לניידות), '
     'צפיפות מוסדות רגישים (בתי חולים, בתי ספר, בתי אבות). '
     'אוכלוסייה עם ניידות מוגבלת ומדד ח"כ נמוך — פגיעה יותר.'),

    ('O — עומס מבצעי (Operational Bottleneck)', '20%',
     'מודד את הקושי בפינוי האוכלוסייה מהאזור. מבוסס על סימולציית ABM (Agent-Based Modeling): '
     'צווארי בקבוק בצירי פינוי, קיבולת כבישים, זמן פינוי ממוצע, ונקודות חסימה צפויות. '
     'אזורים עם מעט צירי פינוי או צפיפות כבישים נמוכה — ציון גבוה יותר.'),

    ('R⁻¹ — גרעון משאבים (Response Capacity, inverted)', '10%',
     'מודד את המחסור במשאבי תגובה. הערך הפוך (inverted): ציון גבוה = פחות משאבים = יותר סיכון. '
     'כולל: מספר ניידות בתחנה, זמן תגובה ממוצע, רמת הכשרה (KYT — Hazard Prediction Training). '
     'תחנה עם מעט ניידות וזמן תגובה ארוך — ציון גבוה.'),

    ('I⁻¹ — חוסן תשתית (Infrastructure, inverted)', '5%',
     'מודד את חוסר עמידות התשתיות. הערך הפוך: ציון גבוה = תשתיות חלשות = יותר סיכון. '
     'כולל: זמינות צירים חלופיים, DTN (Distance to Nearest shelter), מספר מקלטים אנכיים (מבנים 4+ קומות). '
     'אזור ללא מקלטים ועם ציר פינוי יחיד — ציון גבוה.'),
]

for title, weight, desc in components:
    add_rtl_paragraph('')
    add_key_point(f'{title} — משקל {weight}', '')
    add_rtl_paragraph(desc, size=11)

add_rtl_paragraph('')
add_rtl_paragraph('דרגות סיכון:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph('   •  80-100: קריטי (אדום) — נדרשת תגובה מיידית וגיבוי מקסימלי', size=11)
add_rtl_paragraph('   •  60-79: גבוה (כתום) — נדרש שיפור משמעותי בהיערכות', size=11)
add_rtl_paragraph('   •  40-59: בינוני (צהוב) — מצב מספק, נדרש ניטור שוטף', size=11)
add_rtl_paragraph('   •  20-39: נמוך (ירוק) — מוכנות טובה, תחזוקה שוטפת', size=11)
add_rtl_paragraph('   •  0-19: מינימלי (כחול) — מוכנות מצוינת', size=11)

add_separator()

# ================================================
# SLIDE 4: Diagram
# ================================================
add_heading_rtl('שקף 4 — דיאגרמה: קלט → עיבוד → פלט', level=1)
add_heading_rtl('תרשים זרימת המודל הרעיוני', level=2)

add_rtl_paragraph('מטרת השקף:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'שקף מרכזי שמציג את כל תהליך העבודה של המודל בצורה ויזואלית: '
    'אילו נתונים נכנסים, אילו שלבי עיבוד מבוצעים, ומה התוצרים הסופיים — '
    'הן גאוגרפיים (מפות) והן טבלאיים (דוחות).'
)

add_rtl_paragraph('')
add_rtl_paragraph('עמודה 1 — נתוני קלט (7 מקורות):', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
inputs = [
    ('DEM 5m מפ"י', 'מודל גובה דיגיטלי ברזולוציה 5 מטר מהמרכז למיפוי ישראל. משמש לחישוב אזורי הצפה לפי גובה פני השטח.'),
    ('גבולות מוניציפליים — OSM', '554 גבולות רשויות מקומיות מ-OpenStreetMap (admin_level 7+8). מגדירים את הטריטוריה של כל עיר.'),
    ('אוכלוסייה וגילאים — למ"ס', 'נתוני הלשכה המרכזית לסטטיסטיקה: 1,285 ישובים עם פיצול לקבוצות גיל (0-5, 6-18, 19-45, 46-55, 56-64, 65+).'),
    ('אשכול סוציו-אקונומי — למ"ס', 'מדד חברתי-כלכלי (1-10) ל-202 ערים. משמש כ-proxy לניידות, הכנסה, ובעלות רכב.'),
    ('רשת כבישים — OSM', 'כבישים ראשיים ומשניים לחישוב צירי פינוי וזמני נסיעה. כולל סיווג לפי סוג כביש.'),
    ('תחנות משטרה — OSM', 'מיקום תחנות משטרה (amenity=police) מ-OpenStreetMap. מוצגות כנקודות עם אייקון משטרת ישראל.'),
    ('מודל הצפה — GSI', 'מודל הצפות מהמכון הגיאולוגי לפי גובה גל (0.5-10 מטר). מגדיר אזור הצפה ועומק מירבי.'),
]
for title, desc in inputs:
    add_key_point(title, desc)

add_rtl_paragraph('')
add_rtl_paragraph('עמודה 2 — שלבי עיבוד (8 שלבים):', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
steps = [
    'חישוב אזורי הצפה לפי גובה גל נבחר (שכבת buffer מקו החוף)',
    'חיתוך מרחבי (Spatial Intersection) של אזורי ההצפה עם גבולות ערים',
    'חישוב H — אחוז שטח מוצף מתוך שטח העיר הכולל + מרחק מחוף + עומק',
    'חישוב V — שקלול נתוני למ"ס: קשישים, ילדים, אשכול ח"כ, מוסדות רגישים',
    'חישוב O — סימולציית ABM לזמן פינוי ממוצע וזיהוי צווארי בקבוק',
    'חישוב R⁻¹ — הערכת משאבי תגובה: ניידות, זמן תגובה, הכשרה',
    'חישוב I⁻¹ — זיהוי מקלטים אנכיים (4+ קומות), צירים חלופיים, DTN',
    'נרמול Min-Max לסקאלה 0-100 + שקלול משוקלל לציון TSRS סופי',
]
for i, step in enumerate(steps, 1):
    add_rtl_paragraph(f'   שלב {i}:  {step}', size=11)

add_rtl_paragraph('')
add_rtl_paragraph('עמודה 3 — תוצרים:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph('תוצרים גאוגרפיים (מפות):', bold=True, size=12)
add_rtl_paragraph('   •  מפת חום הצפות דינמית — משתנה בזמן אמת לפי סליידר גובה גל', size=11)
add_rtl_paragraph('   •  פוליגוני ערים צבועים לפי ציון TSRS (ירוק → אדום)', size=11)
add_rtl_paragraph('   •  נקודות תחנות משטרה עם אייקון על המפה', size=11)
add_rtl_paragraph('   •  שכבות OSM: כבישים (לפי סיווג), מבנים (לפי עומק הצפה)', size=11)

add_rtl_paragraph('תוצרים טבלאיים (דוחות):', bold=True, size=12)
add_rtl_paragraph('   •  ציון TSRS (0-100) לכל עיר עם דרגת סיכון', size=11)
add_rtl_paragraph('   •  פירוט 5 מדדים (H, V, O, R, I) עם הסבר מילולי', size=11)
add_rtl_paragraph('   •  הנחיות מבצעיות בעברית: ניידות גיבוי, מקלטים, חסימות', size=11)
add_rtl_paragraph('   •  פילוח דמוגרפי: התפלגות גילאים, אשכול ח"כ, הכנסה', size=11)
add_rtl_paragraph('   •  ייצוא PDF לתחנה נבחרת', size=11)

add_separator()

# ================================================
# SLIDE 5: Application
# ================================================
add_heading_rtl('שקף 5 — היישום: GIS אינטראקטיבי', level=1)

add_rtl_paragraph('מטרת השקף:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'להציג את היישום עצמו — מה המשתמש רואה ומה הוא יכול לעשות. '
    'ששת הכרטיסים מתארים את הפיצ\'רים העיקריים של האפליקציה.'
)

add_rtl_paragraph('')
features_detail = [
    ('מפת GIS (Leaflet)', 'מנוע המפה מבוסס Leaflet.js עם שכבות בסיס: OSM + ESRI Hillshade (תבליט). '
     'ניתן להחליף לשכבת לוויין (ESRI), טופוגרפית (OpenTopoMap), או GOVMAP. '
     'המפה תומכת בזום מ-7 עד 18, כולל סקאלה מטרית.'),

    ('גבולות ערים (554)', 'גבולות מוניציפליים אמיתיים של כל הישובים בישראל, מיובאים מ-OpenStreetMap '
     'באמצעות Overpass API. הפוליגונים צבועים לפי ציון TSRS — מירוק (מינימלי) עד אדום (קריטי). '
     'בלחיצה על עיר נפתח פופאפ עם פירוט מלא.'),

    ('סליידר גובה גל', 'המשתמש בוחר גובה גל בין 0.5 ל-10 מטר בצעדים של 0.5. '
     'מפת החום (Heatmap) מתעדכנת בזמן אמת — רדיוס ועוצמה משתנים בהתאם לגובה הגל. '
     'מוצג גם תג חומרה: נמוך / בינוני / חזק / קיצוני.'),

    ('פילוח דמוגרפי', 'בלחיצה על כפתור "פילוח דמוגרפי" בפופאפ, מוצג בסרגל הצד: '
     'גרף עמודות של התפלגות גילאים (5 קבוצות), אשכול סוציו-אקונומי אמיתי מלמ"ס (1-10), '
     'הכנסה ממוצעת ושיעור בעלות רכב — נגזרים מהאשכול.'),

    ('תחנות משטרה (OSM)', 'שכבה נקודתית הנטענת מ-OpenStreetMap בזמן אמת (Overpass API, amenity=police). '
     'כל תחנה מוצגת עם אייקון מגן משטרת ישראל. בלחיצה — שם התחנה, טלפון, כתובת. '
     'השכבה זמינה בזום 12 ומעלה.'),

    ('הנחיות מבצעיות', 'לוח הנחיות דינמי שנפתח בתחתית המסך. כולל: ציון TSRS, '
     'מצב פינוי (חובה/מומלץ/המלצה), ניידות גיבוי נדרשות, מקלטים אנכיים, '
     'מוסדות קריטיים, וצירי פינוי. ההנחיות בעברית ומותאמות לגובה הגל שנבחר.'),
]

for title, desc in features_detail:
    add_rtl_paragraph('')
    add_key_point(title, '')
    add_rtl_paragraph(desc, size=11)

add_separator()

# ================================================
# SLIDE 6: Sources & References
# ================================================
add_heading_rtl('שקף 6 — מקורות מידע וביבליוגרפיה', level=1)

add_rtl_paragraph('מטרת השקף:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'להציג את כל מקורות הנתונים שעליהם מבוסס המודל, ואת המקורות האקדמיים '
    'שמספקים את הבסיס התיאורטי. שקף זה חשוב לביסוס האמינות והלגיטימיות של המודל.'
)

add_rtl_paragraph('')
add_rtl_paragraph('מקורות נתונים (6):', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
data_sources = [
    ('הלמ"ס (הלשכה המרכזית לסטטיסטיקה)', 'נתוני אוכלוסייה, גילאים, אשכול סוציו-אקונומי. 1,285 ישובים + 202 ערים עם מדד ח"כ.'),
    ('מפ"י (המרכז למיפוי ישראל)', 'מודל גובה דיגיטלי (DEM) ברזולוציה 5 מטר. בסיס לחישוב אזורי הצפה.'),
    ('GSI (המכון הגיאולוגי)', 'מודל הצפות צונאמי — מגדיר אזורי הצפה ועומק מירבי לפי גובה גל.'),
    ('OpenStreetMap', 'גבולות מוניציפליים, רשת כבישים, מבנים, תחנות משטרה — נתונים פתוחים.'),
    ('ESRI', 'שכבת World Hillshade (תבליט) — שירות WMS חינמי.'),
    ('GOVMAP', 'שכבות WMS ממשלתיות — מועצות אזוריות, תצלומי אוויר.'),
]
for title, desc in data_sources:
    add_key_point(title, desc)

add_rtl_paragraph('')
add_rtl_paragraph('מקורות ביבליוגרפיים (7):', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
refs = [
    'Wood, N. & Jones, J. (2015). "Vulnerability-Exposure Indices for Tsunami Hazards." USGS Scientific Investigations Report. — בסיס לחישוב מדד H (חשיפה פיזית).',
    'Tucson Police Department. "Workload-Based Staffing Model for Patrol Resource Allocation." — מודל ייחוס להקצאת כוח אדם משטרתי לפי עומס עבודה.',
    'Mas, E. et al. (2012). "Agent-Based Simulation of the 2011 Great East Japan Earthquake Tsunami Evacuation." Natural Hazards. — בסיס לסימולציית ABM (מדד O).',
    'Patel, D. (2019). "GIS-Based Multi-Criteria Analysis for Tsunami Risk Assessment in Coastal Communities." Applied Geography. — מתודולוגיית MCDA לשילוב מדדי סיכון.',
    'Israel Central Bureau of Statistics (2022). "Socio-Economic Profiles of Local Authorities." — מקור נתוני אשכול ח"כ ופרופיל דמוגרפי.',
    'Geist, E. & Parsons, T. (2006). "Probabilistic Analysis of Tsunami Hazards." Natural Hazards. — מסגרת תיאורטית להסתברות צונאמי בים התיכון.',
    'Leone, F. et al. (2011). "A Spatial Analysis of the December 26th, 2004 Tsunami-Induced Damages: Lessons Learned for a Better Risk Assessment." Natural Hazards and Earth System Sciences. — שיטות הערכת נזק מרחבית.',
]
for i, ref in enumerate(refs, 1):
    add_rtl_paragraph(f'   {i}.  {ref}', size=10)

add_separator()

# ================================================
# SLIDE 7: Summary & CTA
# ================================================
add_heading_rtl('שקף 7 — סיכום וקריאה לפעולה', level=1)

add_rtl_paragraph('מטרת השקף:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
add_rtl_paragraph(
    'שקף הסיום מסכם את עיקרי המודל ומציע צעדים קונקרטיים להמשך. '
    'חשוב לסיים בנימה פעולתית — מה הקהל צריך לעשות אחרי המצגת.'
)

add_rtl_paragraph('')
add_rtl_paragraph('נקודות סיכום:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
summary = [
    '5 מדדים משוקללים, ציון TSRS 0-100 — שפה אחידה לתיעדוף ברמה ארצית',
    'חישוב מבוסס GIS + ABM + נתוני למ"ס — לא אינטואיציה',
    'ממשק אינטואיטיבי: סליידר גובה גל + בחירת נפה → תוצאות מיידיות',
    '554 ערים עם גבולות מוניציפליים מלאים — כיסוי ארצי',
    '5 שפות (עברית, אנגלית, רוסית, צרפתית, ספרדית) — נגישות מקסימלית',
    'OSM + Leaflet + Open Source — עלות מינימלית, ללא רישיונות',
]
for s in summary:
    add_rtl_paragraph(f'   ✅  {s}', size=11)

add_rtl_paragraph('')
add_rtl_paragraph('5 הצעדים הבאים:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
cta = [
    ('אישור BRD', 'קבלת אישור מממ"מ / מנהל מחשוב על מסמך הדרישות העסקיות'),
    ('תיאום GSI ומפ"י', 'הסדרת גישה לנתוני הצפה ו-DEM — צפי 4-6 שבועות'),
    ('3 תחנות פיילוט', 'בחירת 3 תחנות חוף ים תיכון להוכחת היתכנות (POC)'),
    ('הקמת תשתית', 'Docker + מסד נתונים — שבוע 1 מתחילת הפיתוח'),
    ('Kick-off', 'ישיבת פתיחה עם צוות GIS + מפקדי תחנות שנבחרו'),
]
for title, desc in cta:
    add_key_point(title, desc)

add_rtl_paragraph('')
add_rtl_paragraph('ציטוט לסיום:', bold=True, color=RGBColor(0x14, 0xB8, 0xA6), size=13)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"מידע גאוגרפי מציל חיים — הכנה מוקדמת מציל יותר"')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
run.font.name = 'Arial'
run.italic = True

# === Save ===
output = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'TSRS_Presenter_Guide.docx')
doc.save(output)
print(f'Presenter guide saved to: {output}')

# Also save to files_2
output2 = r'C:\Users\imark\Desktop\משימת מבחן\files_2\TSRS_Presenter_Guide.docx'
doc.save(output2)
print(f'Also saved to: {output2}')
